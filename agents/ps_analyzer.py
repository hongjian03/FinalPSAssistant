import os
import streamlit as st
import io
from typing import Dict, Any, Optional
import requests
import json
import fitz  # PyMuPDF
import docx
import markitdown
from config.prompts import load_prompts

class PSAnalyzer:
    """
    Agent 2.2: 负责分析用户上传的PS初稿文件，结合院校信息和支持文件分析生成改写策略
    """
    
    def __init__(self, model_name=None):
        """
        初始化PS分析代理。
        
        Args:
            model_name: 使用的LLM模型名称
        """
        # 设置模型名称，如果未提供则使用默认值
        self.model_name = model_name if model_name else "anthropic/claude-3-7-sonnet"
        
        # 从Streamlit secrets获取API密钥
        self.api_key = st.secrets.get("OPENROUTER_API_KEY", "")
        
        # 设置OpenRouter API端点
        self.api_url = "https://openrouter.ai/api/v1/chat/completions"
    
    def analyze_ps(self, ps_file, university_info: str, supporting_file_analysis: str = "未提供支持文件分析", writing_requirements: str = "") -> str:
        """
        分析PS初稿，生成改写策略报告。
        
        Args:
            ps_file: 上传的PS初稿文件
            university_info: 院校信息收集报告
            supporting_file_analysis: 支持文件分析报告
            writing_requirements: 用户的PS写作需求
            
        Returns:
            格式化的PS改写策略报告
        """
        if not ps_file:
            return "未提供PS初稿文件，无法生成改写策略。"
        
        try:
            # 提取PS文件内容
            ps_content = self._extract_ps_content(ps_file)
            
            if not ps_content:
                return "无法从PS文件中提取内容，请检查文件格式。"
            
            # 构建提示
            prompt = self._build_analysis_prompt(ps_content, university_info, supporting_file_analysis, writing_requirements)
            
            # 调用OpenRouter API生成报告
            return self._call_openrouter_api(prompt)
        
        except Exception as e:
            st.error(f"分析PS初稿时出错: {str(e)}")
            return self._get_mock_report()
    
    def _extract_ps_content(self, file) -> str:
        """从上传的PS文件中提取内容"""
        filename = file.name.lower()
        content = ""
        
        try:
            # PDF文件处理
            if filename.endswith(".pdf"):
                # 保存上传的文件
                bytes_data = file.getvalue()
                
                # 使用PyMuPDF读取PDF
                with fitz.open(stream=bytes_data, filetype="pdf") as pdf:
                    for page_num in range(len(pdf)):
                        page = pdf[page_num]
                        content += page.get_text() + "\n\n"
            
            # Word文档处理(.docx)
            elif filename.endswith(".docx"):
                # 使用python-docx读取.docx文件
                doc = docx.Document(io.BytesIO(file.getvalue()))
                content = "\n\n".join([para.text for para in doc.paragraphs if para.text])
                
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_content = " | ".join([cell.text for cell in row.cells])
                        content += f"\n{row_content}"
            
            # Word文档处理(.doc) - 使用markitdown
            elif filename.endswith(".doc"):
                # 使用markitdown尝试提取.doc文件内容
                bytes_data = file.getvalue()
                content = markitdown.convert(bytes_data)
            
            # 其他文本文件处理
            else:
                # 尝试作为文本文件读取
                content = file.getvalue().decode("utf-8")
        
        except Exception as e:
            st.warning(f"无法处理PS文件 {filename}: {str(e)}")
            content = f"[无法处理的文件: {filename}]"
        
        return content
    
    def _build_analysis_prompt(self, ps_content: str, university_info: str, supporting_file_analysis: str, writing_requirements: str = "") -> str:
        """构建分析提示"""
        # 处理写作需求
        writing_req_section = ""
        if writing_requirements and writing_requirements.strip():
            writing_req_section = f"""
        # 用户写作需求
        
        ```
        {writing_requirements}
        ```
        """
        
        # 从配置文件加载提示词
        import os
        import sys
        import json
        
        # 添加父目录到路径以便正确导入
        sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
        from config.prompts import load_prompts
        
        prompts = load_prompts()
        
        # 获取PSAnalyzer的角色、任务和输出格式提示词
        role = prompts["ps_analyzer"]["role"]
        task = prompts["ps_analyzer"]["task"]
        output_format = prompts["ps_analyzer"]["output"]
        
        # 构建完整提示
        prompt = f"""
        # 角色: PS改写策略专家
        
        {role}
        
        # 任务
        
        {task}
        
        # PS初稿内容
        
        ```
        {ps_content}
        ```
        
        # 院校信息收集报告
        
        ```
        {university_info}
        ```
        
        # 支持文件分析报告
        
        ```
        {supporting_file_analysis}
        ```
        {writing_req_section}
        
        # 输出格式
        
        {output_format}
        """
        
        return prompt
    
    def _call_openrouter_api(self, prompt: str) -> str:
        """调用OpenRouter API使用选定的模型生成报告"""
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_key}",
            "HTTP-Referer": "https://ps-assistant.streamlit.app", 
            "X-Title": "PS Assistant Tool"
        }
        
        payload = {
            "model": self.model_name,
            "messages": [{"role": "user", "content": prompt}]
        }
        
        with st.spinner(f"使用 {self.model_name} 分析PS初稿，生成改写策略..."):
            try:
                response = requests.post(self.api_url, headers=headers, json=payload)
                
                if response.status_code == 200:
                    result = response.json()
                    content = result["choices"][0]["message"]["content"]
                    return content
                else:
                    st.error(f"OpenRouter API 错误 ({self.model_name}): {response.status_code} - {response.text}")
                    return self._get_mock_report()
            except Exception as e:
                st.error(f"OpenRouter API 调用错误: {str(e)}")
                return self._get_mock_report()
    
    def _get_mock_report(self) -> str:
        """
        获取模拟报告作为后备选项。
        
        Returns:
            模拟的PS改写策略报告
        """
        mock_report = """
        # PS改写策略报告
        
        ## 整体评估
        PS初稿整体结构清晰，分为学术背景、研究经历、职业目标三大部分。文章逻辑连贯，能够基本说明申请者为什么适合该项目。然而，文章过于陈述性，缺乏深度反思和个人特色；某些段落过于冗长，关键信息被稀释；与目标项目的匹配度有待提高。
        
        ## 与院校匹配分析
        - 初稿对目标院校的具体研究方向和特色提及较少，未能体现针对性
        - 没有充分展示申请者的背景与该项目的具体课程设置或研究重点的匹配度
        - 缺乏对该校特定资源（如实验室、教授、研究中心）的了解和兴趣表达
        - 未能清晰说明为何选择该校该专业，而非其他类似项目
        
        ## 现有优势
        - 学术背景描述较为清晰，体现了扎实的专业基础
        - 研究经历部分具体，包含了技术细节和成果
        - 职业规划明确，表明了长期目标
        - 语言表达基本流畅，没有明显语法错误
        
        ## 需改进方面
        - "告诉"多于"展示"，缺乏具体案例和细节支持核心论点
        - 开头缺乏吸引力，未能立即抓住读者注意力
        - 对研究经历的叙述偏重于描述过程，缺乏对个人贡献和思考的强调
        - 结尾部分薄弱，缺乏有力的总结和对未来的展望
        - 整体缺乏明确的主题线索贯穿全文
        
        ## 支持材料整合建议
        - 将简历中提到的CVPR论文经历详细展开，强调个人贡献和解决问题的过程
        - 整合成绩单中的优秀课程表现，尤其是与目标专业直接相关的课程
        - 利用实习评价中提到的团队协作和解决实际问题的能力，通过具体案例展示
        - 适当提及获奖经历，但需与申请目标和个人发展叙事相结合
        
        ## 段落改写建议
        
        ### 开头段落
        现状：开头直接介绍自己的学术背景，缺乏吸引力。
        建议：以一个引人入胜的研究案例或个人启示作为开场，展示你对该领域的热情和洞察力。然后自然过渡到个人介绍，明确表达申请意向和主题。
        
        ### 学术背景段落
        现状：简单列举了课程和成绩，缺乏深度。
        建议：选择2-3门最相关的关键课程深入描述，说明这些课程如何塑造了你的研究兴趣和专业视角，以及它们与目标项目的关联性。
        
        ### 研究经历段落
        现状：按时间顺序描述了研究过程，但缺乏个人贡献的强调。
        建议：采用STAR法则（情境-任务-行动-结果）重构该部分，强调你面临的挑战、采取的创新方法、个人贡献以及取得的成果。特别突出你的思考过程和解决问题的能力。
        
        ### 实习经历段落
        现状：描述了职责但未能与学术目标建立联系。
        建议：将实习经历与你的研究兴趣和目标院校的研究方向联系起来，说明这些经历如何加深了你对领域问题的理解，以及为何这些经验使你更适合该项目。
        
        ### 目标院校段落
        现状：对目标院校的提及笼统且缺乏具体研究。
        建议：详细说明该院校特定的研究项目、实验室或教授的工作如何与你的兴趣吻合。提及1-2位你希望合作的教授及其研究，表明你已做了充分研究。
        
        ### 结尾段落
        现状：简单重申了申请意向，力度不足。
        建议：有力地总结你的核心优势、与项目的匹配度，以及你能为项目带来什么独特价值。明确表达你的短期学术目标和长期职业规划，展现出清晰的愿景。
        
        ## 改写要点总结
        1. 建立贯穿全文的明确主题，突出你独特的学术发展路径和研究兴趣
        2. 开头增加具体案例或故事元素，提高吸引力
        3. 深入分析而非仅描述研究经历，强调个人贡献和解决问题的能力
        4. 具体研究并提及目标院校的特定项目、教授或研究方向
        5. 将支持材料中的优势点（如CVPR论文、技术技能、实习成果）有机整合到文中
        6. 减少一般性描述，增加具体例子和细节
        7. 强化结尾部分，明确表达短期学术目标和长期职业规划
        8. 整体调整语气，从陈述性转向反思性，展示思考深度和学术成熟度
        """
        
        return mock_report 