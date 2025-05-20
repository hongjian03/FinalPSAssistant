import streamlit as st
import os
from PIL import Image
import io
from datetime import datetime
import uuid
import base64
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import custom modules
from agents.transcript_analyzer import TranscriptAnalyzer
from agents.competitiveness_analyst import CompetitivenessAnalyst
from agents.consulting_assistant import ConsultingAssistant
from agents.serper_client import SerperClient
from config.prompts import load_prompts, save_prompts

# 导入LangSmith追踪功能
from langsmith import traceable

# Set page configuration
st.set_page_config(
    page_title="Applicant Analysis Tool",
    page_icon="🎓",
    layout="wide"
)

# Initialize LangSmith client
def init_langsmith():
    """Initialize LangSmith client from secrets."""
    try:
        langsmith_api_key = st.secrets.get("LANGSMITH_API_KEY", "")
        langsmith_project = st.secrets.get("LANGSMITH_PROJECT", "applicant-analysis-tool")
        
        if langsmith_api_key:
            os.environ["LANGCHAIN_API_KEY"] = langsmith_api_key
            os.environ["LANGCHAIN_PROJECT"] = langsmith_project
            os.environ["LANGCHAIN_TRACING_V2"] = "true"
            os.environ["LANGSMITH_TRACING_V2"] = "true"
            return True
        return False
    except Exception as e:
        st.error(f"Error initializing LangSmith: {str(e)}")
        return False

# 初始化LangSmith
langsmith_enabled = init_langsmith()

# Initialize session state
if "competitiveness_report" not in st.session_state:
    st.session_state.competitiveness_report = None
if "project_recommendations" not in st.session_state:
    st.session_state.project_recommendations = None
if "transcript_content" not in st.session_state:
    st.session_state.transcript_content = None
if "serper_initialized" not in st.session_state:
    st.session_state.serper_initialized = False
if "analyst_model" not in st.session_state:
    st.session_state.analyst_model = "qwen/qwen-max"
if "consultant_model" not in st.session_state:
    st.session_state.consultant_model = "qwen/qwen-max"
if "show_recommendations" not in st.session_state:
    st.session_state.show_recommendations = False
if "transcript_image" not in st.session_state:
    st.session_state.transcript_image = None
if "university" not in st.session_state:
    st.session_state.university = ""
if "major" not in st.session_state:
    st.session_state.major = ""
if "predicted_degree" not in st.session_state:
    st.session_state.predicted_degree = ""
if "custom_requirements" not in st.session_state:
    st.session_state.custom_requirements = ""
if "analysis_status" not in st.session_state:
    st.session_state.analysis_status = None

# Check if necessary API keys are set
def check_api_keys():
    """Check if the necessary API keys are set in Streamlit secrets."""
    api_keys = {
        "OPENROUTER_API_KEY": st.secrets.get("OPENROUTER_API_KEY", None),
        "SERPER_API_KEY": st.secrets.get("SERPER_API_KEY", None),
        "SMITHERY_API_KEY": st.secrets.get("SMITHERY_API_KEY", None),
        "LANGSMITH_API_KEY": st.secrets.get("LANGSMITH_API_KEY", None)
    }
    
    return {k: bool(v) for k, v in api_keys.items()}

# Asynchronously initialize the Serper client
async def init_serper():
    """Initialize the Serper client asynchronously."""
    try:
        # 创建一个主容器用于显示进度和状态，确保显示在主UI而不是侧边栏
        main_container = st.container()
        
        with main_container:
            # 创建进度条标题
            st.subheader("MCP服务初始化")
            
            # 创建专门的进度展示区域，确保进度条靠左对齐
            progress_container = st.container()
            
            with progress_container:
                # 创建进度条和状态文本
                progress_bar = st.progress(0)
                status_text = st.empty()
                status_text.info("正在初始化Serper MCP服务...")
        
            # 检查API密钥
            with progress_container:
                progress_bar.progress(10)
                status_text.info("检查API密钥...")
                
            api_key_status = check_api_keys()
            if not api_key_status.get("SERPER_API_KEY", False) or not api_key_status.get("SMITHERY_API_KEY", False):
                with progress_container:
                    progress_bar.progress(100)
                    status_text.error("缺少必要的API密钥")
                st.error("无法初始化Serper客户端: 缺少必要的API密钥。请确保SERPER_API_KEY和SMITHERY_API_KEY已设置。")
                return False
            
            # 创建新的Serper客户端实例
            with progress_container:
                progress_bar.progress(20)
                status_text.info("创建Serper MCP客户端实例...")
                
            serper_client = SerperClient()
            
            # 尝试初始化，传递主容器以便在其中显示进度
            with progress_container:
                progress_bar.progress(30) 
                status_text.info("开始MCP连接...")
                
            # 让SerperClient的initialize方法处理剩余的进度条更新，传递主容器
            result = await serper_client.initialize(main_container)
            
            if result:
                st.session_state.serper_initialized = True
                st.session_state.serper_client = serper_client  # 保存客户端实例以便重用
                return True
            else:
                st.session_state.serper_initialized = False
                return False
    except Exception as e:
        with main_container:
            st.error(f"初始化Serper客户端时发生异常: {str(e)}")
        st.session_state.serper_initialized = False
        return False

# 支持的模型列表
SUPPORTED_MODELS = [
    "qwen/qwen-max",
    "qwen/qwen3-32b:free",
    "deepseek/deepseek-chat-v3-0324:free",
    "anthropic/claude-3.7-sonnet",
    "openai/gpt-4.1"
]

# 使用LangSmith追踪分析师生成报告的函数
@traceable(run_type="chain", name="CompetitivenessAnalysis")
def generate_competitiveness_report(analyst, university, major, predicted_degree, transcript_content, custom_requirements=""):
    """追踪竞争力分析报告的生成过程"""
    # 直接调用生成报告的方法 - 模型信息会在代理内部记录到 LangSmith
    return analyst.generate_report(
        university=university,
        major=major,
        predicted_degree=predicted_degree,
        transcript_content=transcript_content,
        custom_requirements=custom_requirements
    )

# 使用LangSmith追踪咨询助手推荐项目的函数
@traceable(run_type="chain", name="ProgramRecommendations")
def generate_program_recommendations(consultant, competitiveness_report, custom_requirements=""):
    """追踪项目推荐的生成过程"""
    # 直接调用推荐项目的方法 - 模型信息会在代理内部记录到 LangSmith
    return consultant.recommend_projects(
        competitiveness_report=competitiveness_report,
        custom_requirements=custom_requirements
    )

# 创建Word文档报告并提供下载
def create_downloadable_report(report_title, report_content):
    """生成可下载的Word文档报告"""
    # 创建一个新的Word文档
    doc = docx.Document()
    
    # 设置文档标题
    title = doc.add_heading(report_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加日期
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
    date_run.italic = True
    
    # 添加分隔线
    doc.add_paragraph("_" * 50)
    
    # 添加报告内容 (处理Markdown)
    # 这是一个简单的处理，实际应用中可能需要更复杂的Markdown转换
    lines = report_content.split('\n')
    current_paragraph = None
    
    for line in lines:
        # 处理标题
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # 处理列表项
        elif line.startswith('- ') or line.startswith('* '):
            if not current_paragraph or not current_paragraph.text.startswith(('- ', '* ')):
                current_paragraph = doc.add_paragraph()
            current_paragraph.add_run('\n' + line)
        # 处理空行
        elif not line.strip():
            current_paragraph = None
        # 处理普通文本
        else:
            if not current_paragraph:
                current_paragraph = doc.add_paragraph()
            current_paragraph.add_run(line)
    
    # 保存文档到内存中
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    
    # 转换为Base64编码以便于下载
    base64_docx = base64.b64encode(docx_stream.read()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64_docx}" download="{report_title}.docx">点击下载Word文档</a>'
    
    return href

# Main function
def main():
    # Create tabs
    tab1, tab2, tab3 = st.tabs(["Competitiveness Analysis", "Prompt Debugging", "System Status"])
    
    # 首次加载应用时尝试初始化Serper客户端（仅执行一次）
    if "serper_init_attempted" not in st.session_state:
        st.session_state.serper_initialized = False
        st.session_state.serper_init_attempted = False
    
    # 仅在第一次运行时尝试初始化，避免每次页面刷新都重新连接
    if not st.session_state.serper_init_attempted:
        # 创建主容器用于初始化，确保显示在左侧主UI
        init_container = st.container()
        with init_container:
            st.write("### 初始化网络搜索功能")
            st.write("正在连接到MCP服务，请稍候...")
            
        import asyncio
        asyncio.run(init_serper())
        st.session_state.serper_init_attempted = True
    
    with tab1:
        st.title("Applicant Competitiveness Analysis Tool")
        
        # 检查是否有已生成的报告，决定是否显示表单
        if st.session_state.competitiveness_report is None:
            # 第一阶段：输入基本信息和上传成绩单
            # University selection (currently only one option)
            university = st.selectbox(
                "Select University",
                ["Xi'an Jiaotong-Liverpool University"],
                index=0
            )
            st.session_state.university = university
            
            # Major input
            major = st.text_input("Enter Your Major")
            st.session_state.major = major
            
            # Predicted degree classification
            predicted_degree = st.selectbox(
                "Predicted Degree Classification",
                ["First Class", "Upper Second Class", "Lower Second Class", "Third Class"]
            )
            st.session_state.predicted_degree = predicted_degree
            
            # Transcript upload (可选)
            transcript_file = st.file_uploader(
                "Upload Your Transcript (Optional, Image format only)",
                type=["jpg", "jpeg", "png"]
            )
            
            if transcript_file is not None:
                # 保存图片到会话状态但不显示
                st.session_state.transcript_image = Image.open(transcript_file)
            
            # 添加个性化需求输入框
            custom_requirements = st.text_area(
                "Custom Requirements (Optional)",
                placeholder="Enter any specific requirements or questions you have about UCL programs...",
                help="You can specify particular interests, career goals, or ask specific questions about UCL programs."
            )
            st.session_state.custom_requirements = custom_requirements
            
            # 生成竞争力分析按钮 (仅当必要字段已填写时启用 - 现在只需要专业信息)
            generate_enabled = major and university
            
            # 添加一个信息提示
            if not generate_enabled:
                st.info("Please enter your university and major to generate analysis.")
            
            # 创建右对齐的按钮
            col1, col2, col3 = st.columns([2, 1, 1])
            with col3:
                if st.button("Generate Analysis", disabled=not generate_enabled, key="generate_analysis", use_container_width=True):
                    if major and university:
                        # 从session state获取模型选择和其他信息
                        analyst_model = st.session_state.analyst_model
                        custom_requirements = st.session_state.custom_requirements
                        university = st.session_state.university
                        major = st.session_state.major
                        predicted_degree = st.session_state.predicted_degree
                        
                        # 生成一个会话ID，用于LangSmith追踪
                        session_id = str(uuid.uuid4())
                        
                        # 设置进度状态 - 检查是否需要处理成绩单
                        if st.session_state.transcript_image is not None:
                            st.session_state.analysis_status = "transcript"
                        else:
                            # 如果没有成绩单，直接进入竞争力分析
                            st.session_state.analysis_status = "competitiveness"
                            # 设置空的成绩单内容
                            st.session_state.transcript_content = "No transcript provided."
                        
                        st.rerun()
            
            # 显示分析进度（靠左显示）
            if st.session_state.analysis_status == "transcript":
                # First step: Process the transcript with TranscriptAnalyzer
                with st.spinner("Analyzing transcript with Qwen 2.5 VL via OpenRouter..."):
                    # Process the transcript with AI
                    transcript_analyzer = TranscriptAnalyzer()
                    transcript_content = transcript_analyzer.extract_transcript_data(st.session_state.transcript_image)
                    st.session_state.transcript_content = transcript_content
                
                # 更新状态并重新运行
                st.session_state.analysis_status = "competitiveness"
                st.rerun()
                
            elif st.session_state.analysis_status == "competitiveness":
                # Second step: Generate competitiveness report
                with st.spinner(f"Generating competitiveness report with {st.session_state.analyst_model} via OpenRouter..."):
                    analyst = CompetitivenessAnalyst(model_name=st.session_state.analyst_model)
                    
                    # 使用LangSmith追踪函数包装原始调用
                    if langsmith_enabled:
                        # 使用装饰器追踪的函数
                        with st.status("LangSmith: Tracking competitiveness analysis..."):
                            st.session_state.competitiveness_report = generate_competitiveness_report(
                                analyst,
                                university=st.session_state.university,
                                major=st.session_state.major,
                                predicted_degree=st.session_state.predicted_degree,
                                transcript_content=st.session_state.transcript_content,
                                custom_requirements=st.session_state.custom_requirements
                            )
                    else:
                        # 直接调用函数
                        st.session_state.competitiveness_report = analyst.generate_report(
                            university=st.session_state.university,
                            major=st.session_state.major,
                            predicted_degree=st.session_state.predicted_degree,
                            transcript_content=st.session_state.transcript_content,
                            custom_requirements=st.session_state.custom_requirements
                        )
                
                # 重置状态并重新加载页面以显示结果
                st.session_state.analysis_status = None
                st.rerun()
        
        # 第二阶段：显示结果和推荐按钮
        else:
            # 只有在成绩单存在时才显示成绩单数据
            if st.session_state.transcript_content != "No transcript provided.":
                with st.expander("Transcript Data", expanded=False):
                    st.text_area("Transcript Content", st.session_state.transcript_content, height=200, disabled=True)
            
            # 显示竞争力分析报告（带折叠功能）
            with st.expander("Competitiveness Analysis Report", expanded=True):
                st.markdown(st.session_state.competitiveness_report)
                
                # 添加导出Word文档的按钮
                report_download = create_downloadable_report(
                    "Competitiveness Analysis Report",
                    st.session_state.competitiveness_report
                )
                st.markdown(report_download, unsafe_allow_html=True)
            
            # 显示项目推荐和操作按钮
            col1, col2, col3 = st.columns([2, 1, 1])
            
            with col2:
                # 重置按钮（清除所有结果）
                if st.button("Start Over", key="reset_analysis", use_container_width=True):
                    # 清空所有会话状态
                    st.session_state.competitiveness_report = None
                    st.session_state.project_recommendations = None
                    st.session_state.transcript_content = None
                    st.session_state.transcript_image = None
                    st.session_state.show_recommendations = False
                    st.session_state.custom_requirements = ""
                    st.session_state.analysis_status = None
                    # 重新加载页面
                    st.rerun()
            
            with col3:
                if st.session_state.project_recommendations is None:
                    # 显示推荐按钮（触发项目推荐生成）
                    if st.button("Generate Recommendations", key="generate_recommendations", use_container_width=True):
                        st.session_state.show_recommendations = True
                        
                        # 从session state获取模型选择
                        consultant_model = st.session_state.consultant_model
                        custom_requirements = st.session_state.custom_requirements
                        
                        # 生成项目推荐
                        with st.spinner(f"Generating program recommendations with {consultant_model} via OpenRouter..."):
                            consultant = ConsultingAssistant(model_name=consultant_model)
                            
                            # 使用LangSmith追踪函数包装原始调用
                            if langsmith_enabled:
                                # 使用装饰器追踪的函数
                                with st.status("LangSmith: Tracking program recommendations..."):
                                    st.session_state.project_recommendations = generate_program_recommendations(
                                        consultant,
                                        competitiveness_report=st.session_state.competitiveness_report,
                                        custom_requirements=custom_requirements
                                    )
                            else:
                                # 直接调用函数
                                st.session_state.project_recommendations = consultant.recommend_projects(
                                    competitiveness_report=st.session_state.competitiveness_report,
                                    custom_requirements=custom_requirements
                                )
                            
                            # 重新加载页面以显示结果
                            st.rerun()
            
            # 如果已经生成了项目推荐，则显示
            if st.session_state.project_recommendations is not None:
                with st.expander("UCL Program Recommendations", expanded=True):
                    st.markdown(st.session_state.project_recommendations)
                    
                    # 添加导出Word文档的按钮
                    recommendations_download = create_downloadable_report(
                        "UCL Program Recommendations",
                        st.session_state.project_recommendations
                    )
                    st.markdown(recommendations_download, unsafe_allow_html=True)
    
    with tab2:
        st.title("AI Model & Prompt Configuration")
        
        # 添加模型选择到提示词调试页面顶部
        st.subheader("Model Selection")
        col1, col2 = st.columns(2)
        
        with col1:
            # Model selection for CompetitivenessAnalyst
            analyst_model = st.selectbox(
                "Select Model for Competitiveness Analysis",
                SUPPORTED_MODELS,
                index=SUPPORTED_MODELS.index(st.session_state.analyst_model) if st.session_state.analyst_model in SUPPORTED_MODELS else 0,
                key="analyst_model_debug"
            )
            st.session_state.analyst_model = analyst_model
            
        with col2:
            # Model selection for ConsultingAssistant
            consultant_model = st.selectbox(
                "Select Model for Program Recommendations",
                SUPPORTED_MODELS,
                index=SUPPORTED_MODELS.index(st.session_state.consultant_model) if st.session_state.consultant_model in SUPPORTED_MODELS else 0,
                key="consultant_model_debug"
            )
            st.session_state.consultant_model = consultant_model
        
        # 添加模型选择说明
        st.info("这些模型设置将应用于竞争力分析和项目推荐。您的选择将保存在会话中。")
        
        st.markdown("---")
        
        # Load current prompts
        prompts = load_prompts()
        
        st.subheader("Transcript Analyzer Settings")
        st.markdown("""
        The Transcript Analyzer uses Qwen 2.5 VL (qwen/qwen2.5-vl-72b-instruct) via OpenRouter.
        
        This model is specifically tuned for visual document analysis and transcript data extraction.
        """)
        
        st.subheader("Competitiveness Analyst (Agent 1)")
        
        analyst_role = st.text_area("Role Description", prompts["analyst"]["role"], height=200)
        analyst_task = st.text_area("Task Description", prompts["analyst"]["task"], height=200)
        analyst_output = st.text_area("Output Format", prompts["analyst"]["output"], height=200)
        
        st.subheader("Consulting Assistant (Agent 2)")
        
        consultant_role = st.text_area("Role Description", prompts["consultant"]["role"], height=200)
        consultant_task = st.text_area("Task Description", prompts["consultant"]["task"], height=200)
        consultant_output = st.text_area("Output Format", prompts["consultant"]["output"], height=200)
        
        # 保存按钮，不使用表单
        if st.button("Save Prompts"):
            # Update prompts dictionary
            prompts["analyst"]["role"] = analyst_role
            prompts["analyst"]["task"] = analyst_task
            prompts["analyst"]["output"] = analyst_output
            
            prompts["consultant"]["role"] = consultant_role
            prompts["consultant"]["task"] = consultant_task
            prompts["consultant"]["output"] = consultant_output
            
            # Save updated prompts
            save_prompts(prompts)
            st.success("提示词已成功保存！")

    with tab3:
        st.title("System Status")
        
        # Check API keys
        api_key_status = check_api_keys()
        
        st.subheader("API Keys")
        
        # Display API key status as a table
        status_data = [
            {"API Key": key, "Status": "✅ 已设置" if status else "❌ 未设置"} 
            for key, status in api_key_status.items()
        ]
        
        st.table(status_data)
        
        # LangSmith 状态
        st.subheader("LangSmith Monitoring")
        if langsmith_enabled:
            st.success("✅ LangSmith 监控已启用，两个主要AI代理的输入和输出将被追踪")
            st.info(f"Project: {os.environ.get('LANGCHAIN_PROJECT', 'N/A')}")
            
            # LangSmith 说明
            st.markdown("""
            **LangSmith监控功能：**
            - 追踪竞争力分析和项目推荐的完整请求和响应
            - 记录每个代理的输入参数和输出结果
            - 支持在LangSmith界面上分析和优化提示词
            - 监控模型性能和延迟
            """)
        else:
            st.warning("⚠️ LangSmith 监控未启用。请在secrets中设置 LANGSMITH_API_KEY 以启用此功能")
            
            # 设置说明
            st.markdown("""
            **设置LangSmith：**
            1. 获取LangSmith API密钥: https://smith.langchain.com/
            2. 在`.streamlit/secrets.toml`中添加:
                ```
                LANGSMITH_API_KEY = "your_api_key_here"
                LANGSMITH_PROJECT = "applicant-analysis-tool"  # 可选项
                ```
            """)
        
        # Serper MCP server status
        st.subheader("Serper MCP Server")
        
        # 显示当前状态
        if st.session_state.serper_initialized:
            st.success("✅ Serper MCP客户端已成功初始化，可以进行网络搜索")
        else:
            st.warning("⚠️ Serper MCP客户端未初始化或初始化失败")
        
        # 初始化Serper客户端按钮
        if st.button("重新初始化 Serper MCP客户端"):
            with st.spinner("正在初始化 Serper MCP客户端..."):
                import asyncio
                asyncio.run(init_serper())
                st.rerun()  # 重新加载页面以更新状态
        
        # Add some help text
        st.markdown("""
        ### API 密钥配置
        
        本应用使用 Streamlit secrets 存储 API 密钥。配置 API 密钥的步骤：
        
        1. 创建 `.streamlit/secrets.toml` 文件并添加您的 API 密钥：
           ```toml
           # OpenRouter API (用于访问所有LLM模型，包括视觉模型)
           OPENROUTER_API_KEY = "your_openrouter_api_key"
           
           # Serper Web搜索 API (用于项目推荐)
           SERPER_API_KEY = "your_serper_api_key"
           SMITHERY_API_KEY = "your_smithery_api_key"
           
           # LangSmith监控 API (用于追踪AI代理)
           LANGSMITH_API_KEY = "your_langsmith_api_key"
           LANGSMITH_PROJECT = "applicant-analysis-tool"  # 可选项
           ```
        
        2. 对于 Streamlit Cloud 部署，在 Streamlit Cloud 控制面板中添加这些密钥
        
        ### 常见问题排查
        
        如果遇到MCP连接问题：
        
        1. 确保SERPER_API_KEY和SMITHERY_API_KEY都已正确设置
        2. 检查MCP URL路径是否正确(/mcp而不是/ws)
        3. 确保使用的是streamablehttp_client而不是websocket_client
        4. 如果仍有问题，请查看控制台日志获取详细错误信息
        """)

if __name__ == "__main__":
    main() 